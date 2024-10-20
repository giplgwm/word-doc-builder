import io
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Image as RLImage, Paragraph, PageBreak, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from PIL import Image

MAX_IMG_SIZE = (Inches(6).pt, Inches(8).pt)


def create_word_document(photos, progress_callback=None):
    document = Document()
    last_label = ""

    for i, photo in enumerate(photos):
        if i > 0:
            document.add_page_break()

        with Image.open(photo["path"]) as img:
            with io.BytesIO() as image_stream:
                img.thumbnail(MAX_IMG_SIZE)
                img.save(image_stream, format='JPEG')
                image_stream.seek(0)
                document.add_picture(image_stream)

        label = photo["label"] if photo["label"] else last_label
        last_label = label

        paragraph = document.add_paragraph(label)
        paragraph.alignment = 0

        if progress_callback:
            progress_callback((i + 1) / len(photos))

    doc_file = io.BytesIO()
    document.save(doc_file)
    doc_file.seek(0)

    return doc_file


def create_pdf_document(photos, progress_callback=None):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    last_label = ""

    story = []

    for i, photo in enumerate(photos):
        with Image.open(photo["path"]) as img:
            img_buffer = io.BytesIO()
            img.thumbnail(MAX_IMG_SIZE)
            img.save(img_buffer, format='JPEG')
            img_buffer.seek(0)

            story.append(RLImage(img_buffer))

        label = photo["label"] if photo["label"] else last_label
        last_label = label

        label_style = styles['Normal']
        label_style.alignment = 1  # 1 for center alignment
        story.append(Paragraph(label, label_style))

        if i < len(photos) - 1:
            story.append(PageBreak())

        if progress_callback:
            progress_callback((i + 1) / len(photos))

    doc.build(story)
    buffer.seek(0)

    return buffer
