import streamlit as st
from docx import Document
from io import BytesIO
from PIL import Image

st.set_page_config(
    page_title="Word doc builder",
    page_icon="📝",
    menu_items={'About': "Generates word documents with photos & labels, with page breaks between them."}
)


class ImageCard:
    def __init__(self, img):
        self.img = img
        self.label = img.name.split('.')[0]

    def __str__(self):
        return self.label or "No label provided yet..."

    def change_label(self, new_label: str):
        if (not new_label) or new_label == self.label:
            return
        self.label = new_label


MAX_WIDTH = 512
word_doc_bytesio = None
document_name = ""
image_cards = []
container = st.empty()
files = container.file_uploader("**Labeled Doc Generator**", accept_multiple_files=True, label_visibility='collapsed',
                                help="Upload your images here")


if files:
    container.empty()  # Remove file uploader from page
    with st.form("label_form"):
        # add each file to a doc file
        for user_file in files:
            # Display the files so the user can add their labels
            file_card = ImageCard(user_file)
            image_cards.append(file_card)
            file_card_element = st.empty()
            file_card_text = st.empty()
            file_card_element.image(file_card.img)
            file_card.label = file_card_text.text_input("label1", label_visibility='collapsed', placeholder="Label Photo")
        document_name = st.text_input("Name your word doc", placeholder="Document Name")
        submitted = st.form_submit_button("Generate Word File")

        if submitted:
            # Set up our progress bar
            percentage_per_file = (100 / len(files)) / 100
            current_progress = 0
            percent_text = st.empty()
            percent_text.text("Processing: 0%")
            bar = st.progress(0)

            # Generate doc and BytesIO wrapper
            word_doc = Document()
            word_doc_bytesio = BytesIO()
            for s in image_cards:
                pil_image = Image.open(s.img)
                resized_img = BytesIO()
                if pil_image.width > MAX_WIDTH:
                    aspect_ratio = MAX_WIDTH / pil_image.width
                    new_height = int(pil_image.height * aspect_ratio)
                    pil_image = pil_image.resize((MAX_WIDTH, new_height))
                    pil_image.save(resized_img, format='PNG')
                word_doc.add_picture(resized_img)
                word_doc.add_paragraph(s.label)
                word_doc.add_page_break()

                current_progress += percentage_per_file
                bar.progress(current_progress)
                percent_text.text(f"Processing: {current_progress * 100:.2f}%")
            word_doc.save(word_doc_bytesio)
    if word_doc_bytesio:
        st.download_button("Download Doc file!", word_doc_bytesio, f"{document_name or "Untitled"}.docx")
